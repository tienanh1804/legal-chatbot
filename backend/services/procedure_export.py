"""Export procedure output as DOCX or PDF."""

from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Any, Dict, Optional

logger = logging.getLogger(__name__)


def _fpdf_font_path() -> Optional[str]:
    try:
        import fpdf

        p = Path(fpdf.__file__).resolve().parent / "font" / "DejaVuSans.ttf"
        if p.is_file():
            return str(p)
    except Exception:
        pass
    return None


def export_docx(body: str, out_path: str) -> None:
    """
    Export a pretty DOCX with basic Vietnamese-friendly formatting.

    Heuristics:
    - Recognize common Vietnamese header lines and center/bold them.
    - Use Times New Roman, 13pt, line spacing 1.15.
    - A4 page size, margins ~2cm.
    """
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    import unicodedata
    def _norm_ascii(s: str) -> str:
        """Uppercase + strip Vietnamese diacritics for robust matching."""
        x = (s or "").strip().upper()
        x = x.replace("Đ", "D").replace("đ", "d")
        x = unicodedata.normalize("NFD", x)
        x = "".join(ch for ch in x if unicodedata.category(ch) != "Mn")
        return x


    def _set_run_font(run, size_pt: int = 13, bold: bool = False, underline: bool = False):
        run.bold = bold
        run.underline = underline
        run.font.size = Pt(size_pt)
        # Force Times New Roman for Vietnamese
        run.font.name = "Times New Roman"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")

    def _add_para(doc_obj, text: str, align=None, bold=False, underline=False, size_pt: int = 13):
        p = doc_obj.add_paragraph()
        if align is not None:
            p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.15
        run = p.add_run(text)
        _set_run_font(run, size_pt=size_pt, bold=bold, underline=underline)
        return p

    lines = (body or "").splitlines()
    doc = docx.Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.0)

    sig_title: Optional[str] = None
    pending_sig_labels: Optional[tuple[str, str]] = None
    pending_sig_hints: Optional[tuple[str, str]] = None

    def _add_signature_table(title: Optional[str], labels: tuple[str, str], hints: tuple[str, str]) -> None:
        """Render a 2-column signature block."""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if title:
            _add_para(doc, title, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=False, size_pt=13)

        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Row 1: labels
        for j, text in enumerate(labels):
            cell = table.cell(0, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(2)
            pf.line_spacing = 1.15
            run = p.add_run(text)
            _set_run_font(run, size_pt=13, bold=False)

        # Row 2: hints
        for j, text in enumerate(hints):
            cell = table.cell(1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.15
            run = p.add_run(text)
            _set_run_font(run, size_pt=12, bold=False)

        # Spacer after table
        doc.add_paragraph("")

    for raw in lines:
        line = (raw or "").rstrip()
        if not line.strip():
            # blank line: keep spacing
            doc.add_paragraph("")
            continue

        upper = line.strip().upper()
        norm = _norm_ascii(line)

        if "CONG HOA XA HOI CHU NGHIA VIET NAM" in norm:
            _add_para(
                doc,
                line.strip(),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True,
                size_pt=13,
            )
            continue
        if "DOC LAP" in norm and "TU DO" in norm and "HANH PHUC" in norm:
            _add_para(
                doc,
                line.strip(),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True,
                underline=True,
                size_pt=13,
            )
            continue
        if upper.startswith("---------------") or upper.startswith("—") or upper.startswith("---"):
            _add_para(doc, "---------------", align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size_pt=13)
            continue

        if upper.startswith("DON ") or upper.startswith("ĐƠN "):
            _add_para(doc, line.strip(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=14)
            continue

        if upper.startswith("KINH GUI") or upper.startswith("KÍNH GỬI"):
            # Keep left but emphasize label
            _add_para(doc, line.strip(), align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size_pt=13)
            continue

        # Date/location line before signature is usually right-aligned
        if "NGÀY" in upper or "NGAY" in norm:
            if "……" in line or "..." in line or "….." in line:
                _add_para(doc, line.strip(), align=WD_ALIGN_PARAGRAPH.RIGHT, bold=False, size_pt=13)
                continue

        # Signature block (common in your outputs)
        if "NGƯỜI YÊU CẦU" in upper or "NGUOI YEU CAU" in norm:
            sig_title = line.strip()
            pending_sig_labels = None
            pending_sig_hints = None
            continue

        # Detect "(Vợ) (Chồng)" style line
        if ("(VỢ)" in upper or "(VO)" in norm) and ("(CHỒNG)" in upper or "(CHONG)" in norm):
            pending_sig_labels = ("(Vợ)", "(Chồng)")
            # If there is also a hint embedded on same line, keep as label only
            if pending_sig_hints is not None:
                _add_signature_table(sig_title, pending_sig_labels, pending_sig_hints)
                sig_title = None
                pending_sig_labels = None
                pending_sig_hints = None
            continue

        # Detect "(ký, ghi rõ họ tên)" hints line
        if "(KÝ" in upper or "(KY" in norm:
            # Heuristic: 2 signers in one line
            pending_sig_hints = ("(ký, ghi rõ họ tên)", "(ký, ghi rõ họ tên)")
            if pending_sig_labels is None:
                pending_sig_labels = ("(Vợ)", "(Chồng)")
            _add_signature_table(sig_title, pending_sig_labels, pending_sig_hints)
            sig_title = None
            pending_sig_labels = None
            pending_sig_hints = None
            continue

        _add_para(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size_pt=13)

    doc.save(out_path)


def _replace_placeholders_in_paragraph(paragraph: Any, data: Dict[str, Any]) -> None:
    """
    Replace {{key}} placeholders in a docx paragraph.

    Notes:
    - python-docx splits text into runs; placeholders that are split across runs
      may not be replaced reliably. In practice, typing {{key}} normally usually
      keeps it within one run.
    """
    for run in paragraph.runs:
        txt = run.text or ""
        if "{{" not in txt:
            continue
        for k, v in data.items():
            ph = "{{" + str(k) + "}}"
            if ph in txt:
                txt = txt.replace(ph, str(v))
        run.text = txt


def export_docx_from_template(template_path: str, data: Dict[str, Any], out_path: str) -> None:
    """
    Fill a DOCX template with placeholders and save.

    The template should use placeholders like: {{field_key}} matching the JSON field keys.
    """
    import docx

    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"DOCX template not found: {template_path}")

    safe_data = {str(k): "" if v is None else str(v) for k, v in (data or {}).items()}
    doc = docx.Document(template_path)

    for p in doc.paragraphs:
        _replace_placeholders_in_paragraph(p, safe_data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_placeholders_in_paragraph(p, safe_data)

    doc.save(out_path)


def export_pdf(body: str, out_path: str) -> None:
    from fpdf import FPDF

    font_path = _fpdf_font_path()
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    if font_path:
        pdf.add_font("DejaVu", "", font_path)
        pdf.set_font("DejaVu", size=11)
    else:
        logger.warning("DejaVu font not found; PDF may lack Vietnamese glyphs")
        pdf.set_font("Helvetica", size=11)
    for line in body.split("\n"):
        safe = (
            line.encode("latin-1", errors="replace").decode("latin-1")
            if not font_path
            else line
        )
        pdf.multi_cell(0, 6, txt=safe)
        pdf.ln(2)
    pdf.output(out_path)


def export_procedure(body: str, out_path: str, fmt: str) -> None:
    parent = Path(out_path).parent
    if str(parent) not in ("", "."):
        parent.mkdir(parents=True, exist_ok=True)
    fmt = fmt.lower().strip()
    if fmt == "docx":
        export_docx(body, out_path)
    elif fmt == "pdf":
        export_pdf(body, out_path)
    else:
        raise ValueError(f"Unsupported format: {fmt}")
