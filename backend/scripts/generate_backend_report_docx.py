"""Generate a Vietnamese backend feature report as a .docx file.

This script is intentionally lightweight: it reads no runtime configuration from
the app and does not require the backend server to be running.

Run:
    python backend/scripts/generate_backend_report_docx.py

Output:
    docs/bao_cao_backend.docx
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass(frozen=True)
class FeatureRow:
    """One row in the backend feature summary table."""

    group: str
    feature: str
    endpoints: str
    backend_processing: str
    technologies: str
    notes: Optional[str] = None


def _add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = docx_shared_pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.italic = True


def docx_shared_pt(points: int):
    # Local import keeps dependency surface minimal and avoids type issues.
    from docx.shared import Pt

    return Pt(points)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if h.runs:
        h.runs[0].font.size = docx_shared_pt(14 if level == 1 else 12)


def _add_bullets(doc: Document, items: Iterable[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def _build_rows() -> List[FeatureRow]:
    # Note: "dark mode" (sáng/tối) is typically frontend-only in this repo.
    # We record it as "not handled by backend" to avoid misreporting.
    return [
        FeatureRow(
            group="Xác thực & tài khoản",
            feature="Đăng ký tài khoản",
            endpoints="POST /users/register",
            backend_processing=(
                "Validate email/username, kiểm tra trùng; hash mật khẩu bằng bcrypt; "
                "lưu user vào DB."
            ),
            technologies="FastAPI, Pydantic, SQLAlchemy, passlib[bcrypt]/bcrypt",
        ),
        FeatureRow(
            group="Xác thực & tài khoản",
            feature="Đăng nhập (cấp JWT)",
            endpoints="POST /users/token",
            backend_processing=(
                "Xác thực username/password; tạo JWT access token; trả về token_type=bearer."
            ),
            technologies="FastAPI OAuth2PasswordRequestForm, python-jose[cryptography]",
        ),
        FeatureRow(
            group="Xác thực & tài khoản",
            feature="Xác thực token & thông tin người dùng",
            endpoints="GET /auth/verify; GET /users/me",
            backend_processing="Verify JWT; lấy thông tin user hiện tại (phân quyền theo token).",
            technologies="FastAPI Depends, JWT, middleware/guard (dependency-based)",
        ),
        FeatureRow(
            group="Tài khoản",
            feature="Đổi mật khẩu / Xóa tài khoản",
            endpoints="POST /users/change-password; POST /users/delete-account",
            backend_processing="Verify mật khẩu cũ; cập nhật hash; hoặc xoá user và rollback khi lỗi.",
            technologies="FastAPI, SQLAlchemy transaction, passlib/bcrypt",
        ),
        FeatureRow(
            group="Chatbot RAG",
            feature="Hỏi đáp RAG (có lưu hội thoại khi đăng nhập)",
            endpoints="POST /query",
            backend_processing=(
                "Nhận câu hỏi; (tuỳ chọn) chỉ dùng tài liệu người dùng; chạy RAG (hybrid retrieval + LLM); "
                "format answer/sources; lưu QueryHistory theo conversation_id nếu user đã đăng nhập."
            ),
            technologies=(
                "FastAPI, SQLAlchemy, Hybrid Search (BM25 + vector), FAISS, sentence-transformers, "
                "Google Gemini API"
            ),
            notes="Có hỗ trợ include_user_documents / user_documents_only / user_document_id.",
        ),
        FeatureRow(
            group="Lịch sử hội thoại",
            feature="Lưu và truy xuất lịch sử theo conversation",
            endpoints=(
                "POST /history; GET /history; GET /history/{id}; "
                "GET /history/conversation/{conversation_id}"
            ),
            backend_processing=(
                "Lưu Q/A/sources; tự sinh conversation_id khi chưa có; "
                "truy xuất lịch sử (có chế độ group theo conversation)."
            ),
            technologies="FastAPI, SQLAlchemy, SQLite/PostgreSQL (tuỳ cấu hình)",
        ),
        FeatureRow(
            group="Lịch sử hội thoại",
            feature="Xoá lịch sử",
            endpoints="DELETE /history/{id}; DELETE /history (xoá tất cả)",
            backend_processing="Xoá 1 message hoặc xoá cả conversation; xoá toàn bộ theo user.",
            technologies="FastAPI, SQLAlchemy",
        ),
        FeatureRow(
            group="Tài liệu người dùng (NotebookLM-like)",
            feature="Tải file lên và lập chỉ mục cho RAG theo user",
            endpoints="POST /documents/upload",
            backend_processing=(
                "Nhận multipart upload; giới hạn dung lượng; lưu file theo user; trích xuất text; "
                "chunking; tạo embedding; lưu chunks vào DB để truy hồi theo user."
            ),
            technologies=(
                "FastAPI UploadFile, python-multipart, PyMuPDF (PDF), python-docx (DOCX), Pillow (ảnh), "
                "sentence-transformers, FAISS/Vector search"
            ),
        ),
        FeatureRow(
            group="Tài liệu người dùng",
            feature="Xem file gốc / liệt kê / xoá",
            endpoints="GET /documents/{id}/file; GET /documents; DELETE /documents/{id}",
            backend_processing="Phân quyền theo user; trả file inline; xoá dữ liệu chunks + xoá file vật lý.",
            technologies="FastAPI FileResponse, SQLAlchemy, filesystem",
        ),
        FeatureRow(
            group="Tài liệu người dùng",
            feature="Tóm tắt / trích xuất thông tin từ file",
            endpoints="POST /documents/summarize; POST /documents/extract",
            backend_processing="Ghép text từ chunks đã xử lí; gọi LLM để summarize hoặc extract theo instruction.",
            technologies="Google Gemini API, FastAPI, SQLAlchemy",
        ),
        FeatureRow(
            group="Thủ tục hành chính (Template Wizard)",
            feature="Danh sách template thủ tục",
            endpoints="GET /procedures/templates",
            backend_processing="Đọc các template JSON trong `backend/json_data/procedures` và trả metadata.",
            technologies="FastAPI, JSON",
        ),
        FeatureRow(
            group="Thủ tục hành chính (Template Wizard)",
            feature="Tạo phiên điền thủ tục (wizard) và chat điền từng bước",
            endpoints="POST /procedures/sessions; GET /procedures/sessions/{id}; POST /procedures/sessions/{id}/message",
            backend_processing=(
                "Khởi tạo state; hỏi từng trường theo template; cập nhật state theo câu trả lời; "
                "khi hoàn tất thì render preview_text và chuyển trạng thái completed."
            ),
            technologies="FastAPI, SQLAlchemy, JSON state machine (procedure_service)",
        ),
        FeatureRow(
            group="Thủ tục hành chính (Template Wizard)",
            feature="Tự động prefill từ tài liệu người dùng",
            endpoints="POST /procedures/sessions/{id}/prefill-from-documents",
            backend_processing=(
                "Tổng hợp text từ các file đã upload; dùng LLM trích xuất giá trị cho các field; "
                "điền vào state.collected và quay lại bước hỏi."
            ),
            technologies="Google Gemini API, user-document RAG, FastAPI",
        ),
        FeatureRow(
            group="Thủ tục hành chính (Template Wizard)",
            feature="Xuất thủ tục ra DOCX/PDF",
            endpoints="GET /procedures/sessions/{id}/export?format=docx|pdf",
            backend_processing=(
                "Render nội dung cuối; xuất DOCX (có thể fill theo file mẫu docx_template); "
                "hoặc xuất PDF."
            ),
            technologies="python-docx, fpdf2, FastAPI FileResponse",
        ),
        FeatureRow(
            group="UI sở thích (Sáng/Tối)",
            feature="Dark mode (sáng/tối)",
            endpoints="(frontend xử lí)",
            backend_processing="Không có endpoint lưu theme trong backend hiện tại.",
            technologies="N/A",
            notes="Nếu cần có thể bổ sung `users/settings` để lưu preference.",
        ),
    ]


def _add_table(doc: Document, rows: List[FeatureRow]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Nhóm"
    hdr[1].text = "Chức năng"
    hdr[2].text = "API/Endpoint"
    hdr[3].text = "Backend xử lí"
    hdr[4].text = "Công nghệ"
    hdr[5].text = "Ghi chú"

    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r.group
        cells[1].text = r.feature
        cells[2].text = r.endpoints
        cells[3].text = r.backend_processing
        cells[4].text = r.technologies
        cells[5].text = r.notes or ""


def generate_report(output_path: Path) -> Path:
    """Generate the backend report docx.

    Args:
        output_path: Destination `.docx` path.

    Returns:
        The written output_path.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    _add_title(
        doc,
        title="BÁO CÁO TỔNG HỢP CHỨC NĂNG BACKEND",
        subtitle=f"Dự án: Legal RAG Chatbot — Ngày: {date.today().isoformat()}",
    )

    _add_heading(doc, "1. Tổng quan backend", level=1)
    _add_bullets(
        doc,
        [
            "Backend cung cấp REST API (FastAPI) cho xác thực, chat RAG, lịch sử hội thoại, "
            "upload tài liệu người dùng và module điền thủ tục theo template.",
            "Dữ liệu được quản lí qua SQLAlchemy (migrations bằng Alembic).",
            "Khối tìm kiếm sử dụng hybrid retrieval (BM25 + vector/FAISS) và sinh câu trả lời bằng Gemini.",
        ],
    )

    _add_heading(doc, "2. Công nghệ sử dụng (Backend)", level=1)
    _add_bullets(
        doc,
        [
            "Framework API: FastAPI, Uvicorn; CORS middleware.",
            "Auth: OAuth2 password flow + JWT (python-jose), hash mật khẩu bằng bcrypt (passlib).",
            "Database: SQLAlchemy + Alembic; có thể chạy SQLite/PostgreSQL theo cấu hình.",
            "RAG/Search: sentence-transformers, FAISS, rank-bm25; pipeline hybrid search.",
            "LLM: Google Gemini API (tạo câu trả lời, tóm tắt, trích xuất, prefill field).",
            "Xử lí file: PyMuPDF (PDF), python-docx (DOCX), Pillow (ảnh), giới hạn dung lượng upload.",
            "Xuất thủ tục: DOCX/PDF (python-docx, fpdf2).",
            "Triển khai: Docker, docker-compose.",
        ],
    )

    _add_heading(doc, "3. Bảng tổng hợp chức năng backend", level=1)
    _add_table(doc, _build_rows())

    _add_heading(doc, "4. Ghi chú phạm vi", level=1)
    _add_bullets(
        doc,
        [
            "Chế độ sáng/tối chủ yếu thuộc frontend; backend hiện chưa lưu preference theme.",
            "Các endpoint và luồng xử lí được tổng hợp theo các router trong `backend/api/`.",
        ],
    )

    doc.save(str(output_path))
    return output_path


def main() -> None:
    repo_root = Path(__file__).resolve().parents[2]
    out = repo_root / "docs" / "bao_cao_backend.docx"
    generate_report(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()

